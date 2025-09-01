# src/analysis/run_regressions.py
import os
import argparse
import numpy as np
import pandas as pd
import statsmodels.formula.api as smf
from statsmodels.iolib.summary2 import summary_col

# Optional Word export
try:
    from docx import Document  # python-docx
except Exception:
    Document = None

# ---------- Helpers to make column names robust ----------
ALT_NAMES = {
    "n_A": ["n_A", "n_actionable", "actionable", "nA", "count_actionable", "act_count", "A_count"],
    "n_S": ["n_S", "n_speculative", "speculative", "nS", "count_speculative", "spec_count", "S_count"],
    "n_I": ["n_I", "n_irrelevant", "irrelevant", "nI", "count_irrelevant", "irr_count", "I_count"],
    "n_total": ["n_total", "total_sentences", "doc_count", "sent_count", "total_count"],
    "share_A": ["share_A", "ActShare", "actionable_share", "share_actionable", "AI_frequencyA"],
    "share_S": ["share_S", "SpecShare", "speculative_share", "share_speculative", "AI_frequencyS"],
    "patents_ai": ["patents_ai", "ai_patents", "patent_ai", "patentsAI"],
    "ln_assets": ["ln_assets", "log_assets", "size", "ln_at"],
    "leverage": ["leverage", "lev", "dltt_at", "debt_ratio"],
    "cash": ["cash", "cash_at", "cash_ratio"],
    "rd_intensity": ["rd_intensity", "xrd_at", "rd_at"],
    "capx_at": ["capx_at", "capx_ratio", "capx_over_at"],
    "roa": ["roa", "return_on_assets"],
    "sales_growth": ["sales_growth", "sales_g", "g_sales"],
    "emp": ["emp", "employees", "employment"],
}

# Pretty names for output tables
VAR_LABELS = {
    "n_A": "Actionable (count)",
    "n_S": "Speculative (count)",
    "ActShare": "Actionable share",
    "SpecShare": "Speculative share",
    "log_docs": "log(# AI sentences)",
    "ln_assets": "log(Assets)",
    "leverage": "Leverage",
    "cash": "Cash/Assets",
    "rd_intensity": "R&D/Assets",
    "capx_at": "CAPX/Assets",
    "roa": "ROA",
    "sales_growth": "Sales growth",
    "emp": "Employees (k)",
    "Intercept": "Constant",
    "log_n_A": "log(Actionable mentions + 1)",
    "log_n_S": "log(Speculative mentions + 1)",
    "log_n_I": "log(Irrelevant mentions + 1)",
    "has_actionable": "Actionable disclosure (dummy)",
    "has_spec_only": "Speculative-only disclosure (dummy)",
}

MODEL_TITLES = {
    "OLS_k1_logcounts": "Patents t+1 (log) ~ log mentions (t) + FE",
    "OLS_k1_dummies":   "Patents t+1 (log) ~ disclosure dummies (t) + FE",
    "OLS_k0_logcounts": "Patents t (log) ~ log mentions (t) + FE",
    "LPM_anypat_k1_dummies": "Any AI patent t+1 (LPM) ~ disclosure dummies (t) + FE",
    # Legacy names kept for compatibility if user runs --mode full
    "OLS_k0_levels": "OLS FE (t), levels",
    "OLS_k0_shares": "OLS FE (t), shares",
    "OLS_k2_logcounts": "OLS FE (t+2), log-counts",
    "OLS_k2_dummies": "OLS FE (t+2), dummies",
    "LPM_anypat_k1": "LPM FE (1{patents t+1>0})",
    "LPM_anypat_k1_logcounts": "LPM FE (1{patents t+1>0}), log-counts",
    "OLS_k0_dummies": "OLS FE (t), dummies",
    "OLS_k1_levels": "OLS FE (t+1), levels",
    "OLS_k1_shares": "OLS FE (t+1), shares",
}

REQ_KEYS = ["cik", "year"]

def _resolve(df: pd.DataFrame, key: str):
    for nm in ALT_NAMES[key]:
        if nm in df.columns:
            return nm
    return None

def standardize_columns(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    # counts (if present)
    for k in ["n_A", "n_S", "n_I", "n_total", "patents_ai"]:
        nm = _resolve(df, k)
        if nm and (nm != k):
            df[k] = df[nm]
    # controls
    for k in ["ln_assets","leverage","cash","rd_intensity","capx_at","roa","sales_growth","emp"]:
        nm = _resolve(df, k)
        if nm and (nm != k):
            df[k] = df[nm]
    # shares (if provided already)
    for k in ["share_A","share_S"]:
        nm = _resolve(df, k)
        if nm and (nm != k):
            df[k] = df[nm]
    return df

def _drop_infs(df, cols):
    if not cols:
        return df
    mask = np.isfinite(df[cols].astype(float)).all(axis=1)
    return df.loc[mask].copy()

# ---------- Feature engineering ----------
def add_engineered_cols(df: pd.DataFrame) -> pd.DataFrame:
    df = standardize_columns(df)

    # Ensure required keys exist
    for k in REQ_KEYS:
        if k not in df.columns:
            raise KeyError(f"Required key '{k}' not in panel. Found columns: {list(df.columns)}")

    # log patents
    if "patents_ai" in df.columns:
        df["log_patents_ai"] = np.log1p(df["patents_ai"].fillna(0))

    # Build shares if we have counts
    if all(c in df.columns for c in ["n_A","n_S","n_total"]):
        denom = df["n_total"].replace(0, np.nan)
        df["ActShare"] = df["n_A"] / denom
        df["SpecShare"] = df["n_S"] / denom
        df["log_docs"] = np.log1p(df["n_total"].fillna(0))
    else:
        # Fall back to any provided share columns
        if "share_A" in df.columns:
            df["ActShare"] = df["share_A"]
        if "share_S" in df.columns:
            df["SpecShare"] = df["share_S"]
        # log_docs only if n_total exists
        if "n_total" in df.columns:
            df["log_docs"] = np.log1p(df["n_total"].fillna(0))

    # --- Log-counts (with +1) if counts exist
    if "n_A" in df.columns:
        df["log_n_A"] = np.log1p(df["n_A"].fillna(0).clip(lower=0))
    if "n_S" in df.columns:
        df["log_n_S"] = np.log1p(df["n_S"].fillna(0).clip(lower=0))
    if "n_I" in df.columns:
        df["log_n_I"] = np.log1p(df["n_I"].fillna(0).clip(lower=0))

    # --- Dummies
    if "n_A" in df.columns:
        df["has_actionable"] = (df["n_A"].fillna(0) > 0).astype(int)
    else:
        df["has_actionable"] = np.nan
    if "n_S" in df.columns:
        nA = df["n_A"] if "n_A" in df.columns else 0
        df["has_spec_only"] = ((df["n_S"].fillna(0) > 0) & (pd.Series(nA).fillna(0) == 0)).astype(int)
    else:
        df["has_spec_only"] = np.nan

    # SpecMinusAct only if both shares exist
    if ("SpecShare" in df.columns) and ("ActShare" in df.columns):
        df["SpecMinusAct"] = df["SpecShare"] - df["ActShare"]
    else:
        df["SpecMinusAct"] = np.nan

    return df

def make_leads(df, k_list=(0,1,2)):
    df = df.sort_values(["cik","year"]).copy()
    for k in k_list:
        if "patents_ai" in df.columns:
            df[f"patents_ai_lead{k}"] = df.groupby("cik")["patents_ai"].shift(-k)
            df[f"log_patents_ai_lead{k}"] = np.log1p(df[f"patents_ai_lead{k}"])
            df[f"any_pat_{k}"] = (df[f"patents_ai_lead{k}"].fillna(0) > 0).astype(float)
    return df

def fit_ols_fe(formula, df, cluster="cik", needed=None):
    needed = needed or []
    use = df.dropna(subset=needed).copy()
    use = _drop_infs(use, needed)
    if use.empty:
        raise ValueError("No rows left after dropping NA/inf for required columns.")
    m = smf.ols(formula, data=use).fit(cov_type="cluster", cov_kwds={"groups": use[cluster]})
    return m

def _available_controls(df: pd.DataFrame):
    cand = ["ln_assets","leverage","cash","rd_intensity","capx_at","roa","sales_growth","emp"]
    return [c for c in cand if c in df.columns]

def _star_str(p):
    if p < 0.01:
        return "***"
    if p < 0.05:
        return "**"
    if p < 0.10:
        return "*"
    return ""

def build_clean_table(results_dict):
    """
    Create a compact, journal-friendly table:
    - Rows: selected regressors with pretty labels
    - Columns: each estimated model
    - Cells: coef (se) with significance stars
    Adds bottom notes for N, Firm FE, Year FE.
    """
    # Select the order of terms we want to display
    preferred_order = [
        "n_A", "n_S", "ActShare", "SpecShare",
        "log_n_A", "log_n_S", "log_n_I",
        "has_actionable", "has_spec_only",
        "log_docs",
        "ln_assets", "leverage", "cash", "rd_intensity",
        "capx_at", "roa", "sales_growth", "emp", "Intercept"
    ]
    # Determine which terms actually appear
    present = set()
    for m in results_dict.values():
        present.update(m.params.index.tolist())
    terms = [t for t in preferred_order if t in present]
    # Build a table of formatted coef(se)
    cols = []
    col_names = []
    for key, m in results_dict.items():
        pretty_name = MODEL_TITLES.get(key, key)
        col_names.append(pretty_name)
        col_vals = []
        for t in terms:
            if t in m.params.index:
                coef = m.params[t]
                se = m.bse[t]
                stars = _star_str(m.pvalues[t])
                col_vals.append(f"{coef:.3f}{stars} ({se:.3f})")
            else:
                col_vals.append("")
        cols.append(col_vals)
    # Assemble DataFrame
    import pandas as _pd
    idx = [VAR_LABELS.get(t, t) for t in terms]
    table = _pd.DataFrame({name: col for name, col in zip(col_names, cols)}, index=idx)
    # Add bottom panel
    N_row = {MODEL_TITLES.get(k, k): int(v.nobs) for k, v in results_dict.items()}
    FE_row_firm = {MODEL_TITLES.get(k, k): "Yes" for k in results_dict}
    FE_row_year = {MODEL_TITLES.get(k, k): "Yes" for k in results_dict}
    # Append as rows
    table.loc["N"] = _pd.Series(N_row)
    table.loc["Firm FE"] = _pd.Series(FE_row_firm)
    table.loc["Year FE"] = _pd.Series(FE_row_year)
    return table

def run_all_models(df, outdir, mode="minimal"):
    os.makedirs(outdir, exist_ok=True)
    results = {}

    df = make_leads(df, k_list=(0,1,2))

    have_counts = all(c in df.columns for c in ["n_A","n_S"])
    have_shares = all(c in df.columns for c in ["ActShare","SpecShare"])
    controls = _available_controls(df)
    print(f"[info] Using counts? {have_counts}; Using shares? {have_shares}; Controls: {controls}")

    if not have_counts and not have_shares:
        # print diagnostics to help the user
        print("[error] Neither counts nor shares found. Columns present:", list(df.columns))
        raise RuntimeError("Your panel lacks 'n_A'/'n_S' or 'ActShare'/'SpecShare'. "
                           "If you used prepare_panel_for_regression.py, ensure it wrote ActShare/SpecShare.")

    if mode == "minimal":
        # Ensure any_pat_1 exists
        if "patents_ai_lead1" in df.columns:
            df = df.assign(any_pat_1=(df["patents_ai_lead1"] > 0).astype(float))
        # 1) t+1, log-counts
        if have_counts and all(c in df.columns for c in ["log_n_A","log_n_S"]) and "log_patents_ai_lead1" in df.columns:
            rhs = ["log_n_A","log_n_S"] + (["log_docs"] if "log_docs" in df.columns else []) + controls
            f = f"log_patents_ai_lead1 ~ {' + '.join(rhs)} + C(cik) + C(year)"
            try:
                results["OLS_k1_logcounts"] = fit_ols_fe(f, df, needed=["log_patents_ai_lead1"] + rhs + ["cik","year"])
            except Exception as e:
                print(f"[warn] Minimal: OLS_k1_logcounts failed: {e}")
        # 2) t+1, dummies
        if have_counts and all(c in df.columns for c in ["has_actionable","has_spec_only"]) and "log_patents_ai_lead1" in df.columns:
            rhs = ["has_actionable","has_spec_only"] + (["log_docs"] if "log_docs" in df.columns else []) + controls
            f = f"log_patents_ai_lead1 ~ {' + '.join(rhs)} + C(cik) + C(year)"
            try:
                results["OLS_k1_dummies"] = fit_ols_fe(f, df, needed=["log_patents_ai_lead1"] + rhs + ["cik","year"])
            except Exception as e:
                print(f"[warn] Minimal: OLS_k1_dummies failed: {e}")
        # 3) t, log-counts (contemporaneous)
        if have_counts and all(c in df.columns for c in ["log_n_A","log_n_S"]) and "log_patents_ai_lead0" in df.columns:
            rhs = ["log_n_A","log_n_S"] + (["log_docs"] if "log_docs" in df.columns else []) + controls
            f = f"log_patents_ai_lead0 ~ {' + '.join(rhs)} + C(cik) + C(year)"
            try:
                results["OLS_k0_logcounts"] = fit_ols_fe(f, df, needed=["log_patents_ai_lead0"] + rhs + ["cik","year"])
            except Exception as e:
                print(f"[warn] Minimal: OLS_k0_logcounts failed: {e}")
        # 4) any patent next year, LPM with dummies
        if have_counts and all(c in df.columns for c in ["has_actionable","has_spec_only"]) and "any_pat_1" in df.columns:
            rhs = ["has_actionable","has_spec_only"] + (["log_docs"] if "log_docs" in df.columns else []) + controls
            f = f"any_pat_1 ~ {' + '.join(rhs)} + C(cik) + C(year)"
            try:
                results["LPM_anypat_k1_dummies"] = fit_ols_fe(f, df, needed=["any_pat_1"] + rhs + ["cik","year"])
            except Exception as e:
                print(f"[warn] Minimal: LPM_anypat_k1_dummies failed: {e}")
    else:
        # ---- Baseline OLS FE, k = 0,1,2 on log patents
        for k in [0,1,2]:
            dep = f"log_patents_ai_lead{k}"
            if dep not in df.columns:
                continue
            fe = "+ C(cik) + C(year)"

            # Levels (if counts available)
            if have_counts:
                rhs_terms = ["n_A","n_S"] + (["log_docs"] if "log_docs" in df.columns else []) + controls
                f1 = f"{dep} ~ {' + '.join(rhs_terms)} {fe}"
                try:
                    res1 = fit_ols_fe(f1, df, needed=[dep] + rhs_terms + ["cik","year"])
                    results[f"OLS_k{k}_levels"] = res1
                except Exception as e:
                    print(f"[warn] Levels model k={k} failed: {e}")

            # Shares model if shares available
            if have_shares:
                rhs_terms = ["ActShare","SpecShare"] + (["log_docs"] if "log_docs" in df.columns else []) + controls
                f2 = f"{dep} ~ {' + '.join(rhs_terms)} {fe}"
                try:
                    res2 = fit_ols_fe(f2, df, needed=[dep] + rhs_terms + ["cik","year"])
                    results[f"OLS_k{k}_shares"] = res2
                except Exception as e:
                    print(f"[warn] Shares model k={k} failed: {e}")

            # Log-counts models (only if counts exist)
            if have_counts and all(c in df.columns for c in ["log_n_A","log_n_S"]):
                rhs_terms = ["log_n_A","log_n_S"] + (["log_docs"] if "log_docs" in df.columns else []) + controls
                f_log = f"{dep} ~ {' + '.join(rhs_terms)} {fe}"
                try:
                    res_log = fit_ols_fe(f_log, df, needed=[dep] + rhs_terms + ["cik","year"])
                    results[f"OLS_k{k}_logcounts"] = res_log
                except Exception as e:
                    print(f"[warn] Log-counts model k={k} failed: {e}")

            # Dummies models (only if counts exist)
            if have_counts and all(c in df.columns for c in ["has_actionable","has_spec_only"]):
                rhs_terms = ["has_actionable","has_spec_only"] + (["log_docs"] if "log_docs" in df.columns else []) + controls
                f_dum = f"{dep} ~ {' + '.join(rhs_terms)} {fe}"
                try:
                    res_dum = fit_ols_fe(f_dum, df, needed=[dep] + rhs_terms + ["cik","year"])
                    results[f"OLS_k{k}_dummies"] = res_dum
                except Exception as e:
                    print(f"[warn] Dummies model k={k} failed: {e}")

        # ---- Binary any AI patent (k=1), if patents present
        if "patents_ai_lead1" in df.columns:
            df = df.assign(any_pat_1=(df["patents_ai_lead1"] > 0).astype(float))
            rhs_terms = []
            if have_counts:
                rhs_terms += ["n_A","n_S"]
            if have_shares:
                rhs_terms += ["ActShare","SpecShare"]
            rhs_terms += controls
            if "log_docs" in df.columns:
                rhs_terms += ["log_docs"]
            if rhs_terms:
                f_bin = f"any_pat_1 ~ {' + '.join(rhs_terms)} + C(cik) + C(year)"
                try:
                    res_bin = fit_ols_fe(f_bin, df, needed=["any_pat_1"] + rhs_terms + ["cik","year"])
                    results["LPM_anypat_k1"] = res_bin
                except Exception as e:
                    print(f"[warn] LPM anypat failed: {e}")

            # LPM with log-counts
            if have_counts and all(c in df.columns for c in ["log_n_A","log_n_S"]):
                rhs_terms_log = ["log_n_A","log_n_S"] + controls
                if "log_docs" in df.columns:
                    rhs_terms_log += ["log_docs"]
                f_bin_log = f"any_pat_1 ~ {' + '.join(rhs_terms_log)} + C(cik) + C(year)"
                try:
                    res_bin_log = fit_ols_fe(f_bin_log, df, needed=["any_pat_1"] + rhs_terms_log + ["cik","year"])
                    results["LPM_anypat_k1_logcounts"] = res_bin_log
                except Exception as e:
                    print(f"[warn] LPM anypat (log-counts) failed: {e}")

            # LPM with dummies
            if have_counts and all(c in df.columns for c in ["has_actionable","has_spec_only"]):
                rhs_terms_dum = ["has_actionable","has_spec_only"] + controls
                if "log_docs" in df.columns:
                    rhs_terms_dum += ["log_docs"]
                f_bin_dum = f"any_pat_1 ~ {' + '.join(rhs_terms_dum)} + C(cik) + C(year)"
                try:
                    res_bin_dum = fit_ols_fe(f_bin_dum, df, needed=["any_pat_1"] + rhs_terms_dum + ["cik","year"])
                    results["LPM_anypat_k1_dummies"] = res_bin_dum
                except Exception as e:
                    print(f"[warn] LPM anypat (dummies) failed: {e}")

    if not results:
        raise RuntimeError("No models were estimated. Check that required columns exist and are non-missing.")

    # ---- Save stacked table (TXT + LaTeX) and CSV of coefficients
    if mode == "minimal":
        order_pref = ["OLS_k1_logcounts","OLS_k1_dummies","OLS_k0_logcounts","LPM_anypat_k1_dummies"]
    else:
        order_pref = [
            "OLS_k0_levels","OLS_k0_shares","OLS_k0_logcounts","OLS_k0_dummies",
            "OLS_k1_levels","OLS_k1_shares","OLS_k1_logcounts","OLS_k1_dummies",
            "OLS_k2_levels","OLS_k2_logcounts","OLS_k2_dummies",
            "LPM_anypat_k1","LPM_anypat_k1_logcounts","LPM_anypat_k1_dummies"
        ]
    ordered_keys = [k for k in order_pref if k in results] + [k for k in results.keys() if k not in order_pref]
    ordered = [results[k] for k in ordered_keys]
    readable_names = [MODEL_TITLES.get(k, k) for k in ordered_keys]
    sc = summary_col(
        ordered,
        stars=True,
        float_format="%.3f",
        model_names=readable_names,  # Use readable names here
        info_dict={"N": lambda x: f"{int(x.nobs)}"}
    )
    with open(os.path.join(outdir, "baseline_table.txt"), "w") as f:
        f.write(sc.as_text())
    with open(os.path.join(outdir, "baseline_table.tex"), "w") as f:
        f.write(sc.as_latex())

    rows = []
    for name, m in results.items():
        coefs = m.params.to_frame("coef")
        coefs["se"] = m.bse
        coefs["t"] = m.tvalues
        coefs["p"] = m.pvalues
        coefs["model"] = name
        coefs["N"] = m.nobs
        rows.append(coefs.reset_index().rename(columns={"index": "term"}))
    out = pd.concat(rows, ignore_index=True)
    out.to_csv(os.path.join(outdir, "baseline_coefficients.csv"), index=False)

    # ----- Pretty, journal-style table (Markdown / HTML / optional DOCX)
    clean_tbl = build_clean_table(results)
    md_path = os.path.join(outdir, "baseline_table_clean.md")
    html_path = os.path.join(outdir, "baseline_table_clean.html")
    clean_tbl.to_markdown(md_path)
    clean_tbl.to_html(html_path)

    # Optional Word export if python-docx is available
    if Document is not None:
        try:
            doc = Document()
            doc.add_heading("Baseline regressions", level=1)
            # Add a brief note
            doc.add_paragraph("Entries are coefficients with standard errors in parentheses. *, **, *** denote p<0.10, p<0.05, p<0.01 respectively. All specifications include firm and year fixed effects; SEs clustered by firm.")
            # Add table (simple grid)
            t = doc.add_table(rows=1, cols=1 + len(clean_tbl.columns))
            hdr_cells = t.rows[0].cells
            hdr_cells[0].text = ""
            for j, col in enumerate(clean_tbl.columns, start=1):
                hdr_cells[j].text = str(col)
            for i, (row_name, row) in enumerate(clean_tbl.iterrows()):
                r = t.add_row().cells
                r[0].text = str(row_name)
                for j, col in enumerate(clean_tbl.columns, start=1):
                    r[j].text = str(row[col])
            doc.save(os.path.join(outdir, "baseline_table_clean.docx"))
        except Exception as _e:
            # If anything goes wrong, silently continue after writing md/html
            pass

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--panel", default="data/processed/panel/panel_reg_ready.csv")
    ap.add_argument("--outdir", default="results/01_baseline/tables")
    ap.add_argument("--mode", choices=["minimal","full"], default="minimal",
                    help="minimal = only the 3–4 core models for the paper; full = all variants.")
    args = ap.parse_args()

    df = pd.read_csv(args.panel)

    df = add_engineered_cols(df)

    # Do NOT aggressively drop here; each model drops what's required just-in-time
    run_all_models(df, args.outdir, mode=args.mode)

    print(f"[✓] Wrote clean tables to {args.outdir} (Markdown/HTML and DOCX if available).")

if __name__ == "__main__":
    main()